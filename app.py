import os
from io import BytesIO
from itertools import product

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from scipy.optimize import linprog


# ==================================================
# POMOCNÉ FUNKCE
# ==================================================

def risk_light(prob_loss: float) -> str:
    """Vrátí textový semafor podle pravděpodobnosti ztráty."""
    if prob_loss < 0.05:
        return "🟢 Nízké riziko"
    if prob_loss < 0.15:
        return "🟡 Střední riziko"
    return "🔴 Vysoké riziko"


@st.cache_data
def load_data() -> pd.DataFrame:
    """Načtení demonstračních dat o strojích."""
    return pd.read_csv("data/machines_demo.csv")


# ==================================================
# KONFIGURACE APLIKACE A PŘÍSTUP
# ==================================================

st.set_page_config(page_title="Decision Engine – Půjčovna", layout="centered")

PASSWORD = "sakul7891"
password = st.text_input("Zadejte přístupové heslo", type="password")
if password != PASSWORD:
    st.stop()


# ==================================================
# HLAVIČKA APLIKACE
# ==================================================

st.title("Decision Engine – Půjčovna malostrojů")
st.caption(
    "Interaktivní nástroj pro porovnání investic do stavebních strojů "
    "pomocí simulace nejistoty (Monte Carlo) a optimalizačního modelu (LP, DP)."
)
st.write(
    "Porovnejte dva stroje z hlediska zisku, rizika a návratnosti "
    "při různých scénářích poptávky."
)


# ==================================================
# DATA A SIDEBAR VSTUPY
# ==================================================

df = load_data()
machine_options = df["machine_type"].tolist()

machine_1 = st.sidebar.selectbox("Stroj 1", machine_options, index=0)
machine_2 = st.sidebar.selectbox("Stroj 2", machine_options, index=1)

row_1 = df[df["machine_type"] == machine_1].iloc[0]
row_2 = df[df["machine_type"] == machine_2].iloc[0]

avg_daily_mini = float(row_1["daily_price"])
avg_daily_dumper = float(row_2["daily_price"])

st.sidebar.header("Nastavení simulace")

fixed_cost = st.sidebar.number_input(
    "Fixní náklad (Kč / měsíc)",
    min_value=0,
    value=120000,
    step=5000,
)

scenario = st.sidebar.selectbox(
    "Scénář trhu",
    ["Pesimistický", "Realistický", "Optimistický"],
)

if scenario == "Pesimistický":
    min_days, max_days = 8, 15
    scenario_note = "Slabší poptávka, opatrný scénář."
elif scenario == "Realistický":
    min_days, max_days = 15, 25
    scenario_note = "Běžný provoz firmy."
else:
    min_days, max_days = 22, 30
    scenario_note = "Silná poptávka a vysoké vytížení."

st.sidebar.caption(
    f"{scenario_note} Rozmezí vytížení: {min_days} až {max_days} dní / měsíc."
)

expected_days = st.sidebar.slider("Odhad reálné poptávky (dny/měsíc)", 0, 31, 18)
n = st.sidebar.slider("Počet simulací (měsíců)", 100, 10000, 1000)

st.sidebar.subheader("Optimalizace portfolia")
budget = st.sidebar.number_input(
    "Investiční rozpočet (Kč)",
    min_value=100000,
    value=3000000,
    step=100000,
)


# ==================================================
# MONTE CARLO SIMULACE
# ==================================================

# --- Stroj 1 ---
demand_mean_1 = row_1["demand_days_mean"]
mode_days_1 = np.clip(demand_mean_1, min_days, max_days)
util_days_1 = np.random.triangular(min_days, mode_days_1, max_days, n)

service_cost_1 = np.random.normal(
    row_1["service_cost_mean"],
    row_1["service_cost_std"],
    n,
)
service_cost_1 = np.clip(service_cost_1, 0, None)

failure_happens_1 = np.random.rand(n) < row_1["failure_prob"]
failure_days_1 = np.where(
    failure_happens_1,
    np.random.poisson(row_1["failure_days_mean"], n),
    0,
)
real_days_1 = np.maximum(util_days_1 - failure_days_1, 0)

price_std_1 = avg_daily_mini * 0.08
daily_price_1 = np.random.normal(avg_daily_mini, price_std_1, n)
daily_price_1 = np.clip(daily_price_1, avg_daily_mini * 0.8, avg_daily_mini * 1.2)

rev_mini = real_days_1 * daily_price_1
profit_mini = rev_mini - fixed_cost - service_cost_1

# --- Stroj 2 ---
demand_mean_2 = row_2["demand_days_mean"]
mode_days_2 = np.clip(demand_mean_2, min_days, max_days)
util_days_2 = np.random.triangular(min_days, mode_days_2, max_days, n)

service_cost_2 = np.random.normal(
    row_2["service_cost_mean"],
    row_2["service_cost_std"],
    n,
)
service_cost_2 = np.clip(service_cost_2, 0, None)

failure_happens_2 = np.random.rand(n) < row_2["failure_prob"]
failure_days_2 = np.where(
    failure_happens_2,
    np.random.poisson(row_2["failure_days_mean"], n),
    0,
)
real_days_2 = np.maximum(util_days_2 - failure_days_2, 0)

price_std_2 = avg_daily_dumper * 0.08
daily_price_2 = np.random.normal(avg_daily_dumper, price_std_2, n)
daily_price_2 = np.clip(daily_price_2, avg_daily_dumper * 0.8, avg_daily_dumper * 1.2)

rev_dumper = real_days_2 * daily_price_2
profit_dumper = rev_dumper - fixed_cost - service_cost_2


# ==================================================
# ODVOZENÉ METRIKY A ROZHODOVACÍ LOGIKA
# ==================================================

# --- Realistický scénář při očekávané poptávce ---
expected_profit_mini_real = (
    avg_daily_mini * expected_days - fixed_cost - row_1["service_cost_mean"]
)
expected_profit_dumper_real = (
    avg_daily_dumper * expected_days - fixed_cost - row_2["service_cost_mean"]
)

# --- Základní metriky ---
prob_loss_mini = float(np.mean(profit_mini < 0))
prob_loss_dumper = float(np.mean(profit_dumper < 0))

risk_loss_mini = prob_loss_mini * 100
risk_loss_dumper = prob_loss_dumper * 100

break_even_mini = fixed_cost / avg_daily_mini if avg_daily_mini > 0 else np.nan
break_even_dumper = fixed_cost / avg_daily_dumper if avg_daily_dumper > 0 else np.nan

expected_profit_mini = float(np.mean(profit_mini))
expected_profit_dumper = float(np.mean(profit_dumper))

# --- Percentily / VaR / intervaly ---
p10_mini = float(np.percentile(profit_mini, 10))
p10_dumper = float(np.percentile(profit_dumper, 10))
p50_mini = float(np.percentile(profit_mini, 50))
p50_dumper = float(np.percentile(profit_dumper, 50))
p90_mini = float(np.percentile(profit_mini, 90))
p90_dumper = float(np.percentile(profit_dumper, 90))

var95_mini = float(np.percentile(profit_mini, 5))
var95_dumper = float(np.percentile(profit_dumper, 5))

ci_mini = (p10_mini, p90_mini)
ci_dumper = (p10_dumper, p90_dumper)

prob_dumper_better = float(np.mean(profit_dumper > profit_mini))

# --- ROI ---
investment_cost_mini = row_1["purchase_price"]
investment_cost_dumper = row_2["purchase_price"]

roi_mini = investment_cost_mini / expected_profit_mini if expected_profit_mini > 0 else None
roi_dumper = investment_cost_dumper / expected_profit_dumper if expected_profit_dumper > 0 else None

# --- Executive summary metriky ---
profit_diff = expected_profit_dumper - expected_profit_mini
risk_diff = (prob_loss_mini - prob_loss_dumper) * 100
risk_diff_pp = risk_diff
be_diff = break_even_mini - break_even_dumper
risk_adv_pct = risk_diff

better_machine = machine_2 if expected_profit_dumper > expected_profit_mini else machine_1

if expected_profit_dumper > expected_profit_mini:
    mgmt_text = (
        f"Pro management aktuálně vychází lépe {machine_2}. "
        f"Nabízí vyšší očekávaný zisk, nižší riziko ztráty "
        f"a lepší bod zvratu než {machine_1}. "
        f"Přesto je investice citlivá na cenu pronájmu a vytížení."
    )
elif expected_profit_mini > expected_profit_dumper:
    mgmt_text = (
        f"Pro management aktuálně vychází lépe {machine_1}. "
        f"Nabízí lepší kombinaci očekávaného zisku, rizika "
        f"a návratnosti než {machine_2}."
    )
else:
    mgmt_text = (
        "Obě varianty vycházejí velmi podobně. "
        "Pro finální rozhodnutí doporučujeme doplnit provozní data."
    )

if (prob_loss_dumper < prob_loss_mini) and (break_even_dumper < break_even_mini):
    recommendation = f"✅ Investičně vychází lépe: {machine_2}"
else:
    recommendation = "⚠️ Nelze dát jednoznačné doporučení – doplň data"

if profit_diff > 0:
    profit_text = f"{machine_2} vydělá o {profit_diff:,.0f} Kč více než {machine_1}"
else:
    profit_text = f"{machine_1} vydělá o {abs(profit_diff):,.0f} Kč více než {machine_2}"

if risk_diff > 0:
    risk_text = f"{machine_2} má o {risk_diff:.1f} p.b. nižší riziko ztráty"
else:
    risk_text = f"{machine_1} má o {abs(risk_diff):.1f} p.b. nižší riziko ztráty"


# ==================================================
# OPTIMALIZAČNÍ MODELY – LP A DP
# ==================================================

profit_mini_day = avg_daily_mini - row_1["service_cost_mean"] / 30 - fixed_cost / 30
profit_dumper_day = avg_daily_dumper - row_2["service_cost_mean"] / 30 - fixed_cost / 30

c = [-profit_mini_day, -profit_dumper_day]
A = [[1, 1]]
b = [30]
bounds = [(0, 30), (0, 30)]

lp_result = linprog(c, A_ub=A, b_ub=b, bounds=bounds, method="highs")
if lp_result.success:
    x1_opt, x2_opt = lp_result.x
    lp_profit = -lp_result.fun
else:
    x1_opt, x2_opt = 0, 0
    lp_profit = 0

best_dp_profit = -np.inf
best_d1, best_d2 = 0, 0

for d1 in range(31):
    d2 = 30 - d1
    dp_profit = d1 * profit_mini_day + d2 * profit_dumper_day

    if dp_profit > best_dp_profit:
        best_dp_profit = dp_profit
        best_d1, best_d2 = d1, d2

mc_better = machine_2 if expected_profit_dumper > expected_profit_mini else machine_1
lp_better = (
    machine_1 if x1_opt > x2_opt else machine_2 if x2_opt > x1_opt else "oba stroje podobně"
)
dp_better = (
    machine_1 if best_d1 > best_d2 else machine_2 if best_d2 > best_d1 else "oba stroje podobně"
)


# ==================================================
# SCORE MODELY
# ==================================================

# --- Finální investiční verdikt ---
verdict_score_mini = 0
verdict_score_dumper = 0

if expected_profit_mini > expected_profit_dumper:
    verdict_score_mini += 1
else:
    verdict_score_dumper += 1

if prob_loss_mini < prob_loss_dumper:
    verdict_score_mini += 1
else:
    verdict_score_dumper += 1

if roi_mini is not None and roi_dumper is not None:
    if roi_mini < roi_dumper:
        verdict_score_mini += 1
    else:
        verdict_score_dumper += 1

if x1_opt > x2_opt:
    verdict_score_mini += 1
elif x2_opt > x1_opt:
    verdict_score_dumper += 1

if best_d1 > best_d2:
    verdict_score_mini += 1
elif best_d2 > best_d1:
    verdict_score_dumper += 1

# --- Investiční scorecard ---
scorecard_mini = 0
scorecard_dumper = 0

if expected_profit_mini > expected_profit_dumper:
    scorecard_mini += 1
else:
    scorecard_dumper += 1

if prob_loss_mini < prob_loss_dumper:
    scorecard_mini += 1
else:
    scorecard_dumper += 1

if var95_mini > var95_dumper:
    scorecard_mini += 1
else:
    scorecard_dumper += 1

if roi_mini is not None and roi_dumper is not None:
    if roi_mini < roi_dumper:
        scorecard_mini += 1
    else:
        scorecard_dumper += 1

# --- AI doporučení ---
if scorecard_dumper > scorecard_mini:
    ai_text = f"""
**Proč vychází lépe {machine_2}:**
- Očekávaný měsíční zisk: **{machine_2} {expected_profit_dumper:,.0f} Kč** vs. **{machine_1} {expected_profit_mini:,.0f} Kč**
- {machine_2} má **o {risk_adv_pct:.1f} p.b. nižší riziko ztráty**
- Break-even: **{machine_2} {break_even_dumper:.2f} dní** vs. **{machine_1} {break_even_mini:.2f} dní**

**Scénáře zisku (Kč) – aby bylo jasné riziko:**
- 10% pesimistický scénář: {machine_2} **{p10_dumper:,.0f}**, {machine_1} **{p10_mini:,.0f}**
- Medián: {machine_2} **{p50_dumper:,.0f}**, {machine_1} **{p50_mini:,.0f}**
- 90% optimistický scénář: {machine_2} **{p90_dumper:,.0f}**, {machine_1} **{p90_mini:,.0f}**

**Co je potřeba doplnit před finálním nákupem:**
- reálnou poptávku po {machine_2}
- servisní náklady / odstávky / pojištění
- sezónnost (léto vs. zima)
"""
elif scorecard_mini > scorecard_dumper:
    ai_text = f"""
**Proč vychází lépe {machine_1}:**
- Očekávaný měsíční zisk: **{machine_1} {expected_profit_mini:,.0f} Kč** vs. **{machine_2} {expected_profit_dumper:,.0f} Kč**
- {machine_1} má **o {abs(risk_adv_pct):.1f} p.b. nižší riziko ztráty**
- Break-even: **{machine_1} {break_even_mini:.2f} dní** vs. **{machine_2} {break_even_dumper:.2f} dní**

**Scénáře zisku (Kč) – aby bylo jasné riziko:**
- 10% pesimistický scénář: {machine_1} **{p10_mini:,.0f}**, {machine_2} **{p10_dumper:,.0f}**
- Medián: {machine_1} **{p50_mini:,.0f}**, {machine_2} **{p50_dumper:,.0f}**
- 90% optimistický scénář: {machine_1} **{p90_mini:,.0f}**, {machine_2} **{p90_dumper:,.0f}**

**Co je potřeba doplnit před finálním nákupem:**
- reálnou poptávku po {machine_1}
- servisní náklady / odstávky / pojištění
- sezónnost (léto vs. zima)
"""
else:
    ai_text = f"""
**Proč nejde dát jasné „kup“:**
- Rozdíl není dostatečně jednoznačný při zvolených parametrech.

**Rychlá fakta:**
- Očekávaný zisk {machine_2}: **{expected_profit_dumper:,.0f} Kč**
- Očekávaný zisk {machine_1}: **{expected_profit_mini:,.0f} Kč**
- Riziko ztráty {machine_2}: **{prob_loss_dumper * 100:.1f} %**
- Riziko ztráty {machine_1}: **{prob_loss_mini * 100:.1f} %**

**Další krok:**
Doplň data o poptávce, servisních nákladech a sezónnosti.
"""


# ==================================================
# CITLIVOST A PORTFOLIO – VÝPOČTY
# ==================================================

# --- Citlivost: break-even křivka ---
days_grid = np.arange(0, 32)
profit_curve_mini = avg_daily_mini * days_grid - fixed_cost
profit_curve_dumper = avg_daily_dumper * days_grid - fixed_cost

# --- Citlivost vstupů ---
base_profit_1 = avg_daily_mini * expected_days - fixed_cost - row_1["service_cost_mean"]
base_profit_2 = avg_daily_dumper * expected_days - fixed_cost - row_2["service_cost_mean"]

sensitivity_data = []
drivers = [
    ("Denní cena", "price"),
    ("Vytížení (dny/měsíc)", "days"),
    ("Fixní náklad", "fixed"),
    ("Servisní náklad", "service"),
]

for label, driver in drivers:
    if driver == "price":
        low_1 = (avg_daily_mini * 0.9) * expected_days - fixed_cost - row_1["service_cost_mean"]
        high_1 = (avg_daily_mini * 1.1) * expected_days - fixed_cost - row_1["service_cost_mean"]
        low_2 = (avg_daily_dumper * 0.9) * expected_days - fixed_cost - row_2["service_cost_mean"]
        high_2 = (avg_daily_dumper * 1.1) * expected_days - fixed_cost - row_2["service_cost_mean"]
    elif driver == "days":
        low_1 = avg_daily_mini * (expected_days * 0.9) - fixed_cost - row_1["service_cost_mean"]
        high_1 = avg_daily_mini * (expected_days * 1.1) - fixed_cost - row_1["service_cost_mean"]
        low_2 = avg_daily_dumper * (expected_days * 0.9) - fixed_cost - row_2["service_cost_mean"]
        high_2 = avg_daily_dumper * (expected_days * 1.1) - fixed_cost - row_2["service_cost_mean"]
    elif driver == "fixed":
        low_1 = avg_daily_mini * expected_days - (fixed_cost * 0.9) - row_1["service_cost_mean"]
        high_1 = avg_daily_mini * expected_days - (fixed_cost * 1.1) - row_1["service_cost_mean"]
        low_2 = avg_daily_dumper * expected_days - (fixed_cost * 0.9) - row_2["service_cost_mean"]
        high_2 = avg_daily_dumper * expected_days - (fixed_cost * 1.1) - row_2["service_cost_mean"]
    else:
        low_1 = avg_daily_mini * expected_days - fixed_cost - (row_1["service_cost_mean"] * 0.9)
        high_1 = avg_daily_mini * expected_days - fixed_cost - (row_1["service_cost_mean"] * 1.1)
        low_2 = avg_daily_dumper * expected_days - fixed_cost - (row_2["service_cost_mean"] * 0.9)
        high_2 = avg_daily_dumper * expected_days - fixed_cost - (row_2["service_cost_mean"] * 1.1)

    sensitivity_data.append((f"{label} – {machine_1}", abs(high_1 - low_1)))
    sensitivity_data.append((f"{label} – {machine_2}", abs(high_2 - low_2)))

sensitivity_data = sorted(sensitivity_data, key=lambda x: x[1], reverse=True)
labels = [x[0] for x in sensitivity_data]
values = [x[1] for x in sensitivity_data]

# --- Risk vs Return ---
rr_x1 = prob_loss_mini * 100
rr_x2 = prob_loss_dumper * 100
rr_y1 = expected_profit_mini
rr_y2 = expected_profit_dumper

# --- Heatmapa ---
days_range = np.arange(5, 31, 1)
price_range = np.arange(avg_daily_mini * 0.6, avg_daily_mini * 1.4, avg_daily_mini * 0.05)
profit_matrix = np.zeros((len(price_range), len(days_range)))

for i, price in enumerate(price_range):
    for j, days in enumerate(days_range):
        profit_matrix[i, j] = price * days - fixed_cost

# --- Minimální cena pronájmu ---
required_price_1 = (fixed_cost + row_1["service_cost_mean"]) / max(expected_days, 1)
required_price_2 = (fixed_cost + row_2["service_cost_mean"]) / max(expected_days, 1)

# --- Optimalizace portfolia ---
machine_summary = []
for _, row in df.iterrows():
    expected_unit_profit = (
        row["daily_price"] * row["demand_days_mean"]
        - row["monthly_fixed_cost"]
        - row["service_cost_mean"]
    )
    machine_summary.append(
        {
            "machine_type": row["machine_type"],
            "purchase_price": row["purchase_price"],
            "expected_unit_profit": expected_unit_profit,
        }
    )

machines = sorted(
    machine_summary,
    key=lambda x: x["expected_unit_profit"],
    reverse=True,
)[:5]

best_profit = -1e18
best_combo = None
n_machines = len(machines)
max_units = 2

for counts in product(range(max_units + 1), repeat=n_machines):
    total_cost = sum(counts[i] * machines[i]["purchase_price"] for i in range(n_machines))
    total_profit = sum(counts[i] * machines[i]["expected_unit_profit"] for i in range(n_machines))

    if total_cost <= budget and total_profit > best_profit and sum(counts) > 0:
        best_profit = total_profit
        best_combo = counts

# --- Budget vs profit křivka ---
budget_values = np.arange(500000, 3000001, 500000)
max_profits = []

for test_budget in budget_values:
    best_profit_curve = -1e18

    for counts in product(range(3), repeat=n_machines):
        total_cost = sum(counts[i] * machines[i]["purchase_price"] for i in range(n_machines))
        total_profit = sum(counts[i] * machines[i]["expected_unit_profit"] for i in range(n_machines))

        if total_cost <= test_budget and total_profit > best_profit_curve:
            best_profit_curve = total_profit

    if best_profit_curve == -1e18:
        best_profit_curve = 0

    max_profits.append(best_profit_curve)

# Složka pro export grafů
os.makedirs("outputs", exist_ok=True)


# ==================================================
# EXPORT REPORTU – FUNKCE
# ==================================================

def fig_to_png_bytes(fig) -> bytes:
    img = BytesIO()
    fig.savefig(img, format="png", dpi=200, bbox_inches="tight")
    img.seek(0)
    return img.read()


def build_docx_report() -> bytes:
    doc = Document()

    doc.add_heading("Decision Engine – investiční report", level=1)
    doc.add_paragraph(f"Porovnání: {machine_1} vs. {machine_2}")

    doc.add_heading("Shrnutí výsledků", level=2)
    doc.add_paragraph(f"Fixní náklad (Kč/měsíc): {fixed_cost:,.0f}")
    doc.add_paragraph(f"Rozsah vytížení: {min_days}-{max_days} dní/měsíc")
    doc.add_paragraph(f"Počet simulací: {n:,}")
    doc.add_paragraph(f"Odhad reálné poptávky: {expected_days} dní/měsíc")

    doc.add_heading("Klíčové metriky", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Metrika"
    hdr[1].text = machine_1
    hdr[2].text = machine_2

    rows = [
        ("Průměrná denní cena", f"{avg_daily_mini:,.0f} Kč", f"{avg_daily_dumper:,.0f} Kč"),
        ("Bod zvratu", f"{break_even_mini:.2f} dní", f"{break_even_dumper:.2f} dní"),
        ("Riziko ztráty", f"{risk_loss_mini:.1f} %", f"{risk_loss_dumper:.1f} %"),
        ("Oček. zisk (Monte Carlo)", f"{expected_profit_mini:,.0f} Kč", f"{expected_profit_dumper:,.0f} Kč"),
    ]
    for metric, v1, v2 in rows:
        cells = table.add_row().cells
        cells[0].text = metric
        cells[1].text = v1
        cells[2].text = v2

    doc.add_heading("Distribuce zisku (Monte Carlo)", level=1)
    doc.add_paragraph(
        "Graf ukazuje rozdělení možných scénářů měsíčního zisku na základě Monte Carlo simulace."
    )
    if os.path.exists("outputs/porovnani_zisku.png"):
        doc.add_picture("outputs/porovnani_zisku.png", width=Inches(6))

    doc.add_heading("Rizikové metriky", level=1)
    doc.add_paragraph(
        f"P10 (10 % nejhorších scénářů): {machine_1} = {p10_mini:,.0f} Kč, {machine_2} = {p10_dumper:,.0f} Kč"
    )
    doc.add_paragraph(
        f"Medián zisku: {machine_1} = {p50_mini:,.0f} Kč, {machine_2} = {p50_dumper:,.0f} Kč"
    )
    doc.add_paragraph(
        f"P90 (optimistický scénář): {machine_1} = {p90_mini:,.0f} Kč, {machine_2} = {p90_dumper:,.0f} Kč"
    )

    doc.add_heading("Interpretace výsledků", level=2)
    doc.add_paragraph(
        "Model simuluje možné scénáře poptávky po strojích a zohledňuje náklady na servis, "
        "fixní náklady a pravděpodobnost poruch. Výsledkem je odhad očekávaného měsíčního "
        "zisku a rizika ztráty."
    )
    doc.add_paragraph(
        "Mapa ziskovosti ukazuje, při jaké kombinaci ceny pronájmu a počtu pronajatých dní "
        "se investice dostává do zisku."
    )
    doc.add_paragraph(
        "Optimalizační model testuje možné kombinace nákupu strojů a vybírá portfolio, "
        "které maximalizuje očekávaný měsíční zisk při daném investičním rozpočtu."
    )

    doc.add_heading("Závěr pro management", level=2)
    if expected_profit_mini < 0 and expected_profit_dumper < 0:
        final_report_text = (
            f"Obě investiční varianty jsou při aktuálním nastavení ztrátové. Relativně lépe "
            f"vychází {machine_2}, který dosahuje nižší očekávané ztráty. Před investicí "
            f"doporučujeme upravit cenu pronájmu, zvýšit vytížení nebo snížit fixní náklady."
        )
    elif scorecard_mini > scorecard_dumper:
        final_report_text = (
            f"Na základě simulace vychází jako výhodnější investice {machine_1}. Tento stroj "
            f"nabízí lepší kombinaci očekávaného zisku a rizika."
        )
    elif scorecard_dumper > scorecard_mini:
        final_report_text = (
            f"Na základě simulace vychází jako výhodnější investice {machine_2}. Tento stroj "
            f"nabízí lepší kombinaci očekávaného zisku a rizika."
        )
    else:
        final_report_text = (
            "Obě investiční varianty vycházejí velmi podobně. Doporučujeme doplnit přesnější provozní data."
        )

    doc.add_paragraph(
        f"Doporučujeme zvážit investici do {better_machine}, který vykazuje lepší kombinaci "
        f"očekávaného zisku, nižšího rizika ztráty a nižšího bodu zvratu."
    )

    doc.add_heading("Podmínky ziskovosti", level=1)
    doc.add_paragraph(
        "Investice by byla výrazně atraktivnější při splnění některých z následujících podmínek:"
    )
    doc.add_paragraph(
        "- zvýšení denní ceny pronájmu\n"
        "- vyšší vytížení strojů (více dní pronájmu za měsíc)\n"
        "- snížení fixních nákladů nebo servisních nákladů"
    )
    doc.add_paragraph(final_report_text)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ==================================================
# TABS
# ==================================================

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    ["📊 Shrnutí", "📈 Monte Carlo", "⚠ Riziko", "📉 Citlivost", "🗺 Optimalizace"]
)


# ==================================================
# TAB 1 – SHRNUTÍ
# ==================================================

with tab1:
    st.subheader("Shrnutí pro management")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Očekávaný zisk (Mini bagr)", f"{expected_profit_mini:,.0f} Kč")
    col2.metric("Očekávaný zisk (Pásový dumper)", f"{expected_profit_dumper:,.0f} Kč")
    col3.metric("Pravděpodobnost ztráty (Dumper)", f"{risk_loss_dumper:.1f} %")
    col4.metric("Bod zvratu (Dumper)", f"{break_even_dumper:.2f} dní")

    st.subheader("🧠 Komentář k výsledkům")
    if expected_profit_mini < 0 and expected_profit_dumper < 0:
        st.warning(
            "Při aktuálním nastavení jsou obě investiční varianty ztrátové. "
            "Zvažte zvýšení ceny pronájmu, vyšší vytížení nebo snížení nákladů."
        )
    st.info("Největší vliv na zisk mají denní cena pronájmu a počet dní využití.")

    st.markdown("## 🏆 Doporučená investice")
    col_dec1, col_dec2, col_dec3, col_dec4 = st.columns(4)
    with col_dec1:
        st.metric("Vybraný stroj", better_machine)
    with col_dec2:
        st.metric("Rozdíl v zisku", f"{abs(profit_diff):,.0f} Kč")
    with col_dec3:
        st.metric("Rozdíl v riziku", f"{abs(risk_diff_pp):.1f} p.b.")
    with col_dec4:
        st.metric("Rozdíl v bodu zvratu", f"{abs(be_diff):.2f} dne")

    st.success(f"Investičně vychází lépe: {better_machine}")

    if expected_profit_dumper > expected_profit_mini:
        mgmt_text = (
            f"Pro management aktuálně vychází lépe {machine_2}. "
            f"Nabízí vyšší očekávaný zisk, nižší riziko ztráty "
            f"a lepší bod zvratu než {machine_1}. "
            f"Přesto je investice citlivá na cenu pronájmu a vytížení."
    )

    elif expected_profit_mini > expected_profit_dumper:
        mgmt_text = (
            f"Pro management aktuálně vychází lépe {machine_1}. "
            f"Nabízí lepší kombinaci očekávaného zisku, rizika "
            f"a návratnosti než {machine_2}."
    )

    else:
        mgmt_text = (
            "Obě varianty vycházejí velmi podobně. "
            "Pro finální rozhodnutí doporučujeme doplnit provozní data."
    )
    st.info(mgmt_text)

    st.subheader("📊 Doporučení")
    st.write(recommendation)

    st.subheader("📊 Klíčové rozdíly mezi stroji")
    if profit_diff > 0:
        st.success(profit_text)
    else:
        st.warning(profit_text)

    if risk_diff > 0:
        st.success(risk_text)
    else:
        st.warning(risk_text)

    st.subheader("🧠 Interpretace výsledků")
    st.info(
        "Rozhodnutí vychází z kombinace očekávaného zisku, rizika ztráty a návratnosti investice."
    )

    st.subheader("💰 Návratnost investice (ROI)")
    col_roi1, col_roi2 = st.columns(2)
    with col_roi1:
        if roi_mini is not None:
            st.metric(f"{machine_1}", f"{roi_mini:.1f} měsíců")
        else:
            st.write(f"{machine_1}: nelze spočítat")
    with col_roi2:
        if roi_dumper is not None:
            st.metric(f"{machine_2}", f"{roi_dumper:.1f} měsíců")
        else:
            st.write(f"{machine_2}: nelze spočítat")

    if roi_mini is not None and roi_dumper is not None:
        if roi_dumper < roi_mini:
            st.success(f"{machine_2} se vrátí rychleji než {machine_1}")
        else:
            st.warning(f"{machine_1} se vrátí rychleji než {machine_2}")

    st.subheader("✅ Finální investiční verdikt")
    if verdict_score_mini > verdict_score_dumper:
        st.success(
            f"Celkově vychází lépe **{machine_1}** — má lepší kombinaci zisku, rizika a návratnosti investice."
        )
    elif verdict_score_dumper > verdict_score_mini:
        st.success(
            f"Celkově vychází lépe **{machine_2}** — má lepší kombinaci zisku, rizika a návratnosti investice."
        )
    else:
        st.warning(
            "Obě varianty vycházejí velmi podobně. Pro finální rozhodnutí je vhodné doplnit reálná provozní data."
        )

    st.subheader("🏁 Investiční scorecard")
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        st.metric(f"Score – {machine_1}", f"{scorecard_mini} / 4")
    with col_s2:
        st.metric(f"Score – {machine_2}", f"{scorecard_dumper} / 4")

    if scorecard_mini > scorecard_dumper:
        st.success(f"Scorecard favorizuje **{machine_1}**.")
    elif scorecard_dumper > scorecard_mini:
        st.success(f"Scorecard favorizuje **{machine_2}**.")
    else:
        st.info("Scorecard vychází nerozhodně.")

    st.subheader("Realistický scénář při očekávané poptávce")
    col3, col4 = st.columns(2)
    with col3:
        st.metric(f"{machine_1} 📊 očekávaný zisk", f"{expected_profit_mini_real:,.0f} Kč")
    with col4:
        st.metric(f"{machine_2} 📊 očekávaný zisk", f"{expected_profit_dumper_real:,.0f} Kč")

    if expected_profit_mini_real < 0 and expected_profit_dumper_real < 0:
        st.warning(
            "Při aktuálních parametrech jsou oba stroje ztrátové. "
            "Zvažte vyšší poptávku, nižší fixní náklady nebo vyšší cenu pronájmu."
        )
    elif expected_profit_mini_real < 0 or expected_profit_dumper_real < 0:
        st.info(
            "Jeden ze strojů je při aktuálních parametrech ztrátový. "
            "Doporučuje se porovnat scénáře poptávky a fixních nákladů."
        )

    st.subheader("AI doporučení (vysvětlení v lidské řeči)")
    st.markdown(ai_text)

    st.subheader("Doporučená minimální cena pronájmu")
    col_price1, col_price2 = st.columns(2)
    with col_price1:
        st.metric(f"{machine_1}", f"{required_price_1:,.0f} Kč/den")
    with col_price2:
        st.metric(f"{machine_2}", f"{required_price_2:,.0f} Kč/den")

    st.caption(
        "Minimální cena pronájmu ukazuje, při jaké denní ceně se při zadané poptávce pokryjí fixní a servisní náklady."
    )

    if avg_daily_mini >= required_price_1:
        st.success(f"{machine_1}: aktuální cena pronájmu je dostatečná.")
    else:
        st.warning(f"{machine_1}: aktuální cena pronájmu je pod minimem pro zisk.")

    if avg_daily_dumper >= required_price_2:
        st.success(f"{machine_2}: aktuální cena pronájmu je dostatečná.")
    else:
        st.warning(f"{machine_2}: aktuální cena pronájmu je pod minimem pro zisk.")

    st.subheader("🧾 Interpretace výsledků")
    st.write(
        f"- {machine_1} má riziko ztráty **{risk_loss_mini:.1f} %** při zvolených parametrech.\n"
        f"- {machine_2} má riziko ztráty **{risk_loss_dumper:.1f} %**.\n"
        "Bod zvratu říká, kolik dní v měsíci musí být stroj pronajatý, aby pokryl fixní náklady."
    )

    st.subheader("📄 Export investičního reportu")
    docx_file = build_docx_report()
    st.download_button(
        label="📄 Stáhnout report (DOCX)",
        data=docx_file,
        file_name="decision_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ==================================================
# TAB 2 – MONTE CARLO
# ==================================================

with tab2:
    st.subheader("Klíčové metriky (Monte Carlo)")
    col1, col2 = st.columns(2)

    with col1:
        st.markdown(f"### {machine_1}")
        st.metric("Očekávaný zisk", f"{expected_profit_mini:,.0f} Kč")
        st.metric("Medián zisku", f"{p50_mini:,.0f} Kč")
        st.metric("P10", f"{p10_mini:,.0f} Kč")
        st.metric("P90", f"{p90_mini:,.0f} Kč")
        st.metric("Pravděpodobnost ztráty", f"{risk_loss_mini:.1f} %")

    with col2:
        st.markdown(f"### {machine_2}")
        st.metric("Očekávaný zisk", f"{expected_profit_dumper:,.0f} Kč")
        st.metric("Medián zisku", f"{p50_dumper:,.0f} Kč")
        st.metric("P10", f"{p10_dumper:,.0f} Kč")
        st.metric("P90", f"{p90_dumper:,.0f} Kč")
        st.metric("Pravděpodobnost ztráty", f"{risk_loss_dumper:.1f} %")

    st.subheader("Porovnání distribuce zisku (Monte Carlo)")
    fig_mc = plt.figure()
    plt.hist(profit_mini, bins=25, alpha=0.5, label=machine_1)
    plt.hist(profit_dumper, bins=25, alpha=0.5, label=machine_2)
    plt.axvline(0)
    plt.axvline(p10_mini, linestyle="--", label="p10 mini")
    plt.axvline(p50_mini, linestyle="-", label="median mini")
    plt.axvline(p90_mini, linestyle="--", label="p90 mini")
    plt.axvline(p10_dumper, linestyle="--", label="p10 dumper")
    plt.axvline(p50_dumper, linestyle="-", label="median dumper")
    plt.axvline(p90_dumper, linestyle="--", label="p90 dumper")
    plt.axvline(var95_mini, linestyle=":", label="VaR mini")
    plt.axvline(var95_dumper, linestyle=":", label="VaR dumper")
    plt.xlabel("Zisk / ztráta (Kč)")
    plt.ylabel("Počet simulovaných měsíců")
    plt.title("Distribuce zisku")
    plt.legend()
    fig_mc.savefig("outputs/porovnani_zisku.png", dpi=200, bbox_inches="tight")
    st.pyplot(fig_mc)
    plt.close(fig_mc)

    st.subheader("Kumulativní pravděpodobnost výsledků (CDF)")
    fig_cdf, ax_cdf = plt.subplots(figsize=(10, 6))
    sorted_mini = np.sort(profit_mini)
    sorted_dumper = np.sort(profit_dumper)
    cdf_mini = np.arange(len(sorted_mini)) / len(sorted_mini)
    cdf_dumper = np.arange(len(sorted_dumper)) / len(sorted_dumper)

    ax_cdf.plot(sorted_mini, cdf_mini, label=machine_1)
    ax_cdf.plot(sorted_dumper, cdf_dumper, label=machine_2)
    ax_cdf.axvline(0, color="black", linestyle="--", linewidth=1)
    ax_cdf.text(0, 0.05, "0 Kč (hranice zisku)", rotation=90)
    ax_cdf.text(
        5000,
        0.2,
        f"Pravděpodobnost ztráty\n{machine_1}: {risk_loss_mini:.1f}%\n{machine_2}: {risk_loss_dumper:.1f}%",
        bbox=dict(facecolor="white", alpha=0.7),
    )
    ax_cdf.set_title("Kumulativní pravděpodobnost výsledků")
    ax_cdf.set_xlabel("Zisk / ztráta (Kč)")
    ax_cdf.set_ylabel("Pravděpodobnost")
    ax_cdf.legend()
    ax_cdf.grid(alpha=0.3)
    st.pyplot(fig_cdf)
    plt.close(fig_cdf)

    st.caption(
        "CDF graf ukazuje pravděpodobnost dosažení určitého výsledku nebo horšího. Například bod, kde křivka protíná nulu, ukazuje pravděpodobnost ztráty."
    )

    st.subheader("Interval nejistoty (Monte Carlo)")
    st.write(f"{machine_1}: mezi **{ci_mini[0]:,.0f} Kč** a **{ci_mini[1]:,.0f} Kč** (80 % scénářů)")
    st.write(f"{machine_2}: mezi **{ci_dumper[0]:,.0f} Kč** a **{ci_dumper[1]:,.0f} Kč** (80 % scénářů)")

    st.subheader("Pravděpodobnost lepší investice")
    st.write(
        f"Pravděpodobnost že **{machine_2}** vydělá více než **{machine_1}** je **{prob_dumper_better * 100:.1f} %**."
    )


# ==================================================
# TAB 3 – RIZIKO
# ==================================================

with tab3:
    st.subheader("🚦 Rizikový semafor")
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        st.write(f"*{machine_1}*")
        st.write(risk_light(prob_loss_mini))
    with col_r2:
        st.write(f"*{machine_2}*")
        st.write(risk_light(prob_loss_dumper))

    st.subheader("Provozní riziko")
    st.write(f"{machine_1}: pravděpodobnost poruchy {row_1['failure_prob'] * 100:.1f} %")
    st.write(f"{machine_2}: pravděpodobnost poruchy {row_2['failure_prob'] * 100:.1f} %")
    st.caption("Porucha znamená odstávku stroje a snížení počtu pronajatých dnů.")

    st.subheader("Hodnota v riziku (VaR)")
    st.write(f"{machine_1}: VaR 5 % = **{var95_mini:,.0f} Kč**")
    st.write(f"{machine_2}: VaR 5 % = **{var95_dumper:,.0f} Kč**")
    st.caption("VaR 5 % ukazuje hranici výsledku v nejhorších 5 % simulovaných scénářů.")

    st.subheader("Mapa rizika vs. výnosu")
    fig_rr, ax_rr = plt.subplots()
    ax_rr.scatter(rr_x1, rr_y1, s=120)
    ax_rr.scatter(rr_x2, rr_y2, s=120)
    ax_rr.text(rr_x1 + 0.3, rr_y1, machine_1)
    ax_rr.text(rr_x2 + 0.3, rr_y2, machine_2)
    ax_rr.set_xlabel("Pravděpodobnost ztráty (%)")
    ax_rr.set_ylabel("Očekávaný zisk (Kč)")
    ax_rr.set_title("Riziko vs. výnos")
    ax_rr.grid(True, alpha=0.3)
    ax_rr.axhline(0, linestyle="--")
    st.pyplot(fig_rr)
    plt.close(fig_rr)


# ==================================================
# TAB 4 – CITLIVOST
# ==================================================

with tab4:
    st.subheader("Citlivostní analýza: zisk vs. vytížení (bod zvratu)")
    fig_be, ax_be = plt.subplots()
    ax_be.plot(days_grid, profit_curve_mini, label=machine_1)
    ax_be.plot(days_grid, profit_curve_dumper, label=machine_2)
    ax_be.axhline(0)
    ax_be.axvline(break_even_mini, linestyle="--")
    ax_be.axvline(break_even_dumper, linestyle="--")
    ax_be.set_xlabel("Vytížení (dny/měsíc)")
    ax_be.set_ylabel("Zisk / ztráta (Kč)")
    ax_be.set_title("Zisk podle vytížení (fixed_cost = zvolený fixní náklad)")
    ax_be.legend()
    st.pyplot(fig_be)
    plt.close(fig_be)

    st.caption(
        f"Bod zvratu {machine_1} = {break_even_mini:.2f} dní, {machine_2} = {break_even_dumper:.2f} dní\n"
        "Nad těmito hodnotami je stroj v průměru v zisku."
    )

    st.subheader("Citlivost zisku na změnu vstupů")
    fig_tornado, ax_tornado = plt.subplots(figsize=(10, 6))
    ax_tornado.barh(labels, values, alpha=0.8)
    ax_tornado.invert_yaxis()
    ax_tornado.set_title("Citlivost zisku na změnu vstupů")
    ax_tornado.set_xlabel("Dopad na očekávaný zisk (Kč)")
    ax_tornado.set_ylabel("Parametr")
    ax_tornado.grid(axis="x", linestyle="--", alpha=0.4)
    st.pyplot(fig_tornado)
    plt.close(fig_tornado)

    st.caption(
        "Tornado graf ukazuje, které vstupy mají největší vliv na očekávaný zisk. Čím delší pruh, tím citlivější je výsledek na změnu daného parametru."
    )

    st.subheader("Heatmapa ziskovosti (poptávka × cena pronájmu)")
    fig_heat, ax_heat = plt.subplots()
    cbar_data = ax_heat.imshow(
        profit_matrix,
        aspect="auto",
        origin="lower",
        extent=[days_range.min(), days_range.max(), price_range.min(), price_range.max()],
    )
    ax_heat.contour(
        days_range,
        price_range,
        profit_matrix,
        levels=[0],
        linewidths=2,
        linestyles="--",
    )
    if profit_matrix.max() > 0:
        ax_heat.contourf(
            days_range,
            price_range,
            profit_matrix,
            levels=[0, profit_matrix.max()],
            colors=["green"],
            alpha=0.3,
        )
    ax_heat.set_xlabel("Poptávka (dny/měsíc)")
    ax_heat.set_ylabel("Cena pronájmu (Kč/den)")
    ax_heat.set_title("Zisk podle ceny a poptávky")
    fig_heat.colorbar(cbar_data, ax=ax_heat, label="Zisk (Kč)")
    st.pyplot(fig_heat)
    plt.close(fig_heat)

    if profit_matrix.max() <= 0:
        st.warning(
            "Při aktuálních parametrech není v zadaném rozsahu ceny a poptávky žádná zisková kombinace."
        )
        st.caption(
            "Pro dosažení zisku je nutné zvýšit cenu pronájmu nebo dosáhnout vyššího vytížení stroje."
        )

    st.caption(
        "Heatmapa ukazuje, při jaké kombinaci ceny pronájmu a počtu pronajatých dní se investice dostává do zisku."
    )
    st.caption(
        "Barevná škála ukazuje očekávaný měsíční zisk. Tmavé barvy znamenají ztrátu, světle zelené až žluté oblasti vysoký zisk."
    )
    st.caption(
        "Přerušovaná čára označuje hranici break-even, tedy kombinace ceny a poptávky, kde je očekávaný zisk roven nule."
    )


# ==================================================
# TAB 5 – OPTIMALIZACE
# ==================================================

with tab5:
    st.header("Optimalizační modely")
    st.caption(
        "Sekce ukazuje výsledky optimalizačních modelů (LP a DP) a následně optimální portfolio strojů při zadaném rozpočtu."
    )

    st.subheader("📐 Deterministická optimalizace (LP)")
    col_lp1, col_lp2, col_lp3 = st.columns(3)
    with col_lp1:
        st.metric(f"Optimální dny – {machine_1}", f"{x1_opt:.1f}")
    with col_lp2:
        st.metric(f"Optimální dny – {machine_2}", f"{x2_opt:.1f}")
    with col_lp3:
        st.metric("Max. deterministický zisk", f"{lp_profit:,.0f} Kč")

    if x1_opt > x2_opt:
        st.info(f"Deterministický model preferuje více využívat {machine_1}.")
    elif x2_opt > x1_opt:
        st.info(f"Deterministický model preferuje více využívat {machine_2}.")
    else:
        st.info("Deterministický model rozděluje využití mezi oba stroje podobně.")

    st.subheader("🧠 Diskrétní optimalizace (DP)")
    col_dp1, col_dp2, col_dp3 = st.columns(3)
    with col_dp1:
        st.metric(f"Optimální dny – {machine_1}", f"{best_d1}")
    with col_dp2:
        st.metric(f"Optimální dny – {machine_2}", f"{best_d2}")
    with col_dp3:
        st.metric("Max. diskrétní zisk", f"{best_dp_profit:,.0f} Kč")

    if best_d1 > best_d2:
        st.info(f"DP model preferuje více využívat {machine_1}.")
    elif best_d2 > best_d1:
        st.info(f"DP model preferuje více využívat {machine_2}.")
    else:
        st.info("DP model rozděluje využití mezi oba stroje podobně.")

    st.subheader("🧠 Porovnání rozhodovacích modelů")
    st.write("**Monte Carlo** pracuje s nejistotou a ukazuje rozdělení možných výsledků.")
    st.write("**LP optimalizace** hledá nejlepší řešení při pevných deterministických předpokladech.")
    st.info(
        f"Monte Carlo preferuje **{mc_better}**.\n\n"
        f"LP preferuje **{lp_better}**.\n\n"
        f"DP preferuje **{dp_better}**."
    )

    st.subheader("📦 Optimalizace nákupu strojů")
    st.caption(
        "Optimalizační model testuje možné kombinace nákupu strojů a vybírá portfolio, které maximalizuje očekávaný měsíční zisk při daném investičním rozpočtu."
    )

    if best_combo is not None:
        st.success(f"Nejlepší portfolio při rozpočtu {budget:,.0f} Kč")
        for i, count in enumerate(best_combo):
            if count > 0:
                st.write(f"{machines[i]['machine_type']}: **{count} ks**")

        total_cost = sum(best_combo[i] * machines[i]["purchase_price"] for i in range(n_machines))
        st.write(f"Celková investice: **{total_cost:,.0f} Kč**")
        st.write(f"Očekávaný měsíční zisk portfolia: **{best_profit:,.0f} Kč**")
    else:
        st.warning("Při daném rozpočtu se nepodařilo najít vhodnou kombinaci.")

    st.caption(
        "Optimalizace vybírá nejlepší kombinaci strojů z celého datasetu při zadaném investičním rozpočtu."
    )

    st.subheader("📈 Rozpočet vs maximální zisk portfolia")
    fig_budget, ax_budget = plt.subplots()
    ax_budget.plot(budget_values, max_profits)
    ax_budget.set_xlabel("Investiční rozpočet (Kč)")
    ax_budget.set_ylabel("Maximální měsíční zisk portfolia")
    ax_budget.set_title("Rozpočet vs zisk portfolia")
    st.pyplot(fig_budget)
    plt.close(fig_budget)

    st.caption(
        "Graf ukazuje maximální očekávaný měsíční zisk portfolia při různých investičních rozpočtech."
    )
