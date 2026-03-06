from docx import Document

import pandas as pd

import numpy as np

import matplotlib.pyplot as plt

print("Decision Engine běží 🚀")

# Načtení Excel souboru
df = pd.read_excel("data/mvp_pujcovna_malostroju.xlsx")

print("Data načtena úspěšně.")
print(df.head())

print("\n--- INFO O DATECH ---")
print(df.info())

print("\n--- ZÁKLADNÍ STATISTIKA ---")
print(df.describe())

print("\nPočet řádků:", len(df))

# 1) Počet sloupců
print("\n--- POČET SLOUPCŮ ---")
print("Počet sloupců:", df.shape[1])
print("Názvy sloupců:", list(df.columns))

# 2) Průměr denní ceny podle typu
print("\n--- PRŮMĚRNÁ DENNÍ CENA PODLE TYPU ---")
avg_daily_by_type = df.groupby("type")["daily_cost_czk"].mean().sort_values(ascending=False)
print(avg_daily_by_type)

# 3) Medián hodinové ceny
print("\n--- MEDIÁN HODINOVÉ CENY (hourly_cost_czk) ---")
median_hourly = df["hourly_cost_czk"].median()
print("Medián hodinové ceny:", median_hourly)

# 4) Nejdražší stroj podle denní ceny
print("\n--- NEJDRAŽŠÍ STROJ PODLE DENNÍ CENY ---")
idx = df["daily_cost_czk"].idxmax()
most_expensive_machine = df.loc[idx]
print(most_expensive_machine)

df = pd.read_excel("data/mvp_pujcovna_malostroju.xlsx")

print("\n--- MINI BAGRY ---")

mini_bagry = df[df["type"] == "Mini Bagr"]

print(mini_bagry)
print("Počet Mini Bagrů:", len(mini_bagry))

print("\n--- NEJDRAŽŠÍ MINI BAGR ---")

idx = mini_bagry["daily_cost_czk"].idxmax()
nejdrazsi_mini_bagr = mini_bagry.loc[idx]

print(nejdrazsi_mini_bagr)

print("\n--- MINI BAGR: DETERMINISTICKÉ VYTÍŽENÍ (20 DNÍ / MĚSÍC) ---")

# 1) průměrná denní cena Mini Bagrů
avg_daily_mini = mini_bagry["daily_cost_czk"].mean()

# 2) deterministické vytížení
util_days = 20

# 3) měsíční tržba (hrubý odhad)
monthly_revenue_mini = avg_daily_mini * util_days

print("Průměrná denní cena Mini Bagrů:", round(avg_daily_mini, 2))
print("Vytížení (dní):", util_days)
print("Odhad měsíčních tržeb 1 Mini Bagru:", round(monthly_revenue_mini, 0), "Kč")

print("\n--- MINI BAGR: SIMULACE VYTÍŽENÍ (15–25 DNÍ) ---")

n = 1000
util_days_samples = np.random.randint(15, 26, size=n)  # 15 až 25 včetně
monthly_revenue_samples = avg_daily_mini * util_days_samples

print("Simulace:", n, "měsíců")
print("Min tržba:", round(monthly_revenue_samples.min(), 0), "Kč")
print("Průměr tržby:", round(monthly_revenue_samples.mean(), 0), "Kč")
print("Max tržba:", round(monthly_revenue_samples.max(), 0), "Kč")

print("10% percentil:", round(np.percentile(monthly_revenue_samples, 10), 0), "Kč")
print("50% (medián):", round(np.percentile(monthly_revenue_samples, 50), 0), "Kč")
print("90% percentil:", round(np.percentile(monthly_revenue_samples, 90), 0), "Kč")

print("\n--- SIMULACE ZISKU (FIXNÍ NÁKLAD 120 000 Kč) ---")

fixed_cost = 120000

profit_samples = monthly_revenue_samples - fixed_cost

print("Min zisk:", round(profit_samples.min(), 0), "Kč")
print("Průměrný zisk:", round(profit_samples.mean(), 0), "Kč")
print("Max zisk:", round(profit_samples.max(), 0), "Kč")

# pravděpodobnost ztráty
prob_loss = np.mean(profit_samples < 0)

print("Pravděpodobnost ztráty:", round(prob_loss * 100, 2), "%")

print("\n--- BREAK-EVEN VYTÍŽENÍ ---")
break_even_days = fixed_cost / avg_daily_mini
print("Break-even (dní/měsíc):", round(break_even_days, 2))

print("\n--- PASOVÝ DUMPER ---")

dumper = df[df["type"] == "Pásový dumper"]

avg_daily_dumper = dumper["daily_cost_czk"].mean()

print("Průměrná denní cena Dumper:", round(avg_daily_dumper, 2))

print("\n--- SIMULACE DUMPER (15–25 DNÍ) ---")

n = 1000
util_days_samples_dumper = np.random.randint(15, 26, size=n)
monthly_revenue_dumper = avg_daily_dumper * util_days_samples_dumper

print("Min tržba:", round(monthly_revenue_dumper.min(), 0))
print("Průměr tržby:", round(monthly_revenue_dumper.mean(), 0))
print("Max tržba:", round(monthly_revenue_dumper.max(), 0))

profit_dumper = monthly_revenue_dumper - fixed_cost
prob_loss_dumper = np.mean(profit_dumper < 0)

print("Pravděpodobnost ztráty Dumper:", round(prob_loss_dumper * 100, 2), "%")

print("\n--- BREAK-EVEN DUMPER ---")
break_even_dumper = fixed_cost / avg_daily_dumper
print("Break-even (dní/měsíc):", round(break_even_dumper, 2))

print("\n--- CITLIVOST FIXNÍCH NÁKLADŮ (DUMPER) ---")

test_fixed_costs = [120000, 140000, 160000, 180000, 200000]

for cost in test_fixed_costs:
    profit_test = monthly_revenue_dumper - cost
    prob_loss_test = np.mean(profit_test < 0)
    print("Fixní náklad:", cost,
          "| Pravděpodobnost ztráty:",
          round(prob_loss_test * 100, 2), "%")
    
    print("\n--- GRAF DISTRIBUCE TRŽEB MINI BAGR ---")

plt.hist(monthly_revenue_samples, bins=20)
plt.title("Distribuce měsíčních tržeb - Mini Bagr")
plt.xlabel("Měsíční tržba (Kč)")
plt.ylabel("Počet měsíců")

plt.show()

print("\n--- GRAF DISTRIBUCE ZISKU MINI BAGR ---")

plt.hist(profit_samples, bins=20)
plt.title("Distribuce měsíčního zisku - Mini Bagr")
plt.xlabel("Zisk / ztráta (Kč)")
plt.ylabel("Počet měsíců")

plt.axvline(0)  # hranice ztráty

plt.show()

print("\n--- POROVNÁNÍ ZISKU: MINI BAGR vs DUMPER ---")

plt.hist(profit_samples, bins=20, alpha=0.5, label="Mini Bagr")
plt.hist(profit_dumper, bins=20, alpha=0.5, label="Dumper")

plt.axvline(0)

plt.title("Porovnání distribuce zisku")
plt.xlabel("Zisk / ztráta (Kč)")
plt.ylabel("Počet měsíců")

plt.legend()

plt.savefig("outputs/porovnani_zisku.png", dpi=200, bbox_inches="tight")
plt.close()

# --- AUTOMATICKÝ REPORT ---
fixed_cost = 120000  # nastavíme zpět základní scénář pro rozhodnutí

# Mini Bagr (použij existující profit_samples / prob_loss / break-even, nebo přepočítej)
prob_loss_mini = np.mean(profit_samples < 0)
avg_daily_mini = mini_bagry["daily_cost_czk"].mean()
break_even_mini = fixed_cost / avg_daily_mini

# Dumper (použij existující profit_dumper)
prob_loss_dumper = np.mean(profit_dumper < 0)
avg_daily_dumper = dumper["daily_cost_czk"].mean()
break_even_dumper = fixed_cost / avg_daily_dumper

# Doporučení
if (prob_loss_dumper < prob_loss_mini) and (break_even_dumper < break_even_mini):
    recommendation = "KUP DUMPER"
else:
    recommendation = "NEKUPOVAT / POTŘEBUJEME DALŠÍ DATA"

report = f"""
AUTOMATICKÝ REPORT – PŮJČOVNA MALOSTROJŮ
========================================

PŘEDPOKLADY
- Fixní náklad: {fixed_cost:,} Kč / měsíc
- Vytížení simulace: 15–25 dní / měsíc
- Počet simulovaných měsíců: {len(profit_samples)}

MINI BAGR
- Průměrná denní cena: {avg_daily_mini:,.1f} Kč
- Break-even: {break_even_mini:.2f} dní / měsíc
- Pravděpodobnost ztráty: {prob_loss_mini*100:.1f} %

PÁSOVÝ DUMPER
- Průměrná denní cena: {avg_daily_dumper:,.1f} Kč
- Break-even: {break_even_dumper:.2f} dní / měsíc
- Pravděpodobnost ztráty: {prob_loss_dumper*100:.1f} %

DOPORUČENÍ
- {recommendation}

VÝSTUPY
- Graf: outputs/porovnani_zisku.png
"""

with open("outputs/report.txt", "w", encoding="utf-8") as f:
    f.write(report)

print("Hotovo: outputs/report.txt a outputs/porovnani_zisku.png")
print("Doporučení:", recommendation)

# --- WORD REPORT ---

doc = Document()

doc.add_heading("Analýza investice – Půjčovna malostrojů", 0)

doc.add_heading("Shrnutí", level=1)

doc.add_paragraph(
    f"Simulace ukazuje, že pásový dumper generuje stabilnější zisk než mini bagr. "
    f"Pravděpodobnost ztráty u mini bagru je {prob_loss_mini*100:.1f} %, "
    f"zatímco u dumperu {prob_loss_dumper*100:.1f} %."
)

doc.add_heading("Výsledky simulace", level=1)

doc.add_paragraph(f"Průměrná denní cena Mini Bagr: {avg_daily_mini:.0f} Kč")
doc.add_paragraph(f"Průměrná denní cena Dumper: {avg_daily_dumper:.0f} Kč")

doc.add_paragraph(f"Break-even Mini Bagr: {break_even_mini:.2f} dní")
doc.add_paragraph(f"Break-even Dumper: {break_even_dumper:.2f} dní")

doc.add_heading("Doporučení", level=1)

doc.add_paragraph(recommendation)

doc.add_heading("Vizualizace", level=1)

doc.add_picture("outputs/porovnani_zisku.png")

doc.save("outputs/report.docx")

print("Word report vytvořen: outputs/report.docx")

import os
os.makedirs("outputs", exist_ok=True)

plt.savefig("outputs/porovnani_zisku.png", dpi=200, bbox_inches="tight")
plt.close()